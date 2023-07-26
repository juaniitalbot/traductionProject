import pandas as pd
import streamlit as st
from docx import Document
import re
import mysql.connector
from itertools import combinations
import base64

def separate_sentences(text):
    # Utilizamos una expresión regular para separar el texto en oraciones.
    sentences = re.split(r'(?<!\w\.\w.)(?<![A-Z][a-z]\.)(?<=\.|\?|\!)\s', text)  # Modified regex to handle exclamation marks as well

    return sentences

def remove_punctuation(word):
    # Eliminar signos de puntuación del inicio y final de la palabra
    return re.sub(r'^\W+|\W+$', '', word)

def main():
    st.set_page_config(layout="wide")  # Set the app to display in wide mode

    st.title("Aplicación de Streamlit para separar oraciones y buscar palabras")
    st.markdown("---")

    # Configurar la conexión a MySQL (modifica los detalles según tu caso)
    host = 'talbot.com.ar'
    user = 'talbot_tradusr'
    password = '3jV5tDTT'
    database = 'talbot_tradprj'
    conn = mysql.connector.connect(host=host, user=user, password=password, database=database)
    c = conn.cursor()

    # Cargar la base de datos de palabras en español e inglés
    word_data = {}  # Diccionario para almacenar los datos de palabras
    c.execute('SELECT spanish, english FROM glosario')
    for row in c.fetchall():
        spanish, english = row
        word_data[english.lower()] = spanish  # Convertimos las palabras en inglés a minúsculas como clave para buscar en el diccionario

    uploaded_file = st.file_uploader("Cargar archivo de Word", type=["docx"])
    st.markdown("---")

    # Columnas
    col1, col2 = st.columns([1, 1])

    if uploaded_file is not None:
        doc = Document(uploaded_file)
        full_text = ""

        for para in doc.paragraphs:
            full_text += para.text + " "

        sentences = separate_sentences(full_text)

        
        for i, sentence in enumerate(sentences):
            col1.write(sentence)

            # Dividir la oración en palabras y eliminar signos de puntuación
            words = re.findall(r'\b\w+\b', sentence)

            # Crear una lista con todas las opciones para el selector
            options = []
            for j, word in enumerate(words):
                word_cleaned = remove_punctuation(word)  # Eliminar signos de puntuación
                if word_cleaned.lower() in word_data:
                    options.append(word_cleaned)
                if j < len(words) - 1:
                    combined_word = " ".join(words[j:j + 2])
                    combined_word_cleaned = remove_punctuation(combined_word)  # Eliminar signos de puntuación de la combinación
                    if combined_word_cleaned.lower() in word_data:
                        options.append(combined_word_cleaned)

            # Selector para elegir una palabra o combinación de palabras de la oración
            selected_option = col1.selectbox("Selecciona una palabra o combinación de palabras:", options)
            
            # Buscar coincidencias en la base de datos y mostrar las traducciones
            translations = []
            if selected_option.lower() in word_data:
                translations.append(word_data[selected_option.lower()])

            col1.success("Traducciones encontradas:")
            for translation in translations:
                col1.write(translation)
            col1.markdown('----')

    # Cerramos la conexión a la base de datos
    conn.close()

    # Estilo para mejorar la apariencia visual de la columna de edición
    col2.markdown('<style>div.row-widget.stRadio > div{flex-direction:column;} </style>', unsafe_allow_html=True)

    edited_text = col2.text_area("Campo de traducción: ", height=500)  # Ajustar la altura del campo de texto

    # Botón para guardar como Word
    if col2.button("Guardar como Word"):
        doc = Document()
        doc.add_paragraph(edited_text)
        col2.markdown(get_binary_file_downloader_html(doc, "documento_editado.docx"), unsafe_allow_html=True)

def get_binary_file_downloader_html(bin_file, file_label='File'):
    with open(bin_file, 'rb') as f:
        data = f.read()
    bin_str = base64.b64encode(data).decode()
    href = f'<a href="data:application/octet-stream;base64,{bin_str}" download="{file_label}.docx">Descargar {file_label}</a>'
    return href

if __name__ == "__main__":
    main()
